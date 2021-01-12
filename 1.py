from docx.api import Document
from xlsxwriter import Workbook
import pandas as pd 
from pandas import DataFrame
import io
import Document_Automation_Program.folder_finding

#import Workbook

path= "C:/Users/ygkwo/Desktop/test1/1.docx" #input_path
output_path="C:/Users/ygkwo/Desktop/test1" #output_path


#original

# document=Document(path)

# writer=pd.ExcelWriter('{}/docx_talbes.xlsx'.format(output_path),engine='xlsxwriter')



# for i in range(len(document.tables)):
#     table=document.tables[i]
#     data=[]
#     keys=None
#     row_data=None
    
#     for j,row in enumerate(table.rows):
#         text=(cell.text for cell in row.cells)
#         if j==0:
#             keys=tuple(text)
#             continue
#         row_data=dict(zip(keys,text))
#         data.append(row_data)
#         #
#         # df1=pd.DataFrame(data)
#         # print(df1)
#         # df1.to_excel(writer,sheet_name='table_{}'.format(i))

#     df=pd.DataFrame(data)
#     print(df)
#     df.to_excel(writer,sheet_name='table_{}'.format(i))


# writer.save()







#docx_to_xlsx Function [Input, Output]
def docx_to_xlsx(input,output):
    document=Document(path)

    writer=pd.ExcelWriter('{}/docx_talbes.xlsx'.format(output),engine='xlsxwriter')

    for i in range(len(document.tables)):
        table=document.tables[i]
        data=[]
        keys=None
        row_data=None
        
        for j,row in enumerate(table.rows):
            text=(cell.text for cell in row.cells)
            if j==0:
                keys=tuple(text)
                continue
            row_data=dict(zip(keys,text))
            data.append(row_data)
            #
            # df1=pd.DataFrame(data)
            # print(df1)
            # df1.to_excel(writer,sheet_name='table_{}'.format(i))

        df=pd.DataFrame(data)
        print(df)
        df.to_excel(writer,sheet_name='table_{}'.format(i))


    writer.save()



def docx_to_xlsx_multi(input,output):
    document=Document(path)

    writer=pd.ExcelWriter('{}/docx_talbes.xlsx'.format(output),engine='xlsxwriter')

    for i in range(len(document.tables)):
        table=document.tables[i]
        data=[]
        keys=None
        row_data=None
        
        for j,row in enumerate(table.rows):
            text=(cell.text for cell in row.cells)
            if j==0:
                keys=tuple(text)
                continue
            row_data=dict(zip(keys,text))
            data.append(row_data)
            #
            # df1=pd.DataFrame(data)
            # print(df1)
            # df1.to_excel(writer,sheet_name='table_{}'.format(i))

        df=pd.DataFrame(data)
        print(df)
        df.to_excel(writer,sheet_name='table_{}'.format(i))


    writer.save()



#one docx file
docx_to_xlsx(path,output_path)
