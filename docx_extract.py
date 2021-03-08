from docx.api import Document
from xlsxwriter import Workbook
import pandas as pd 
from pandas import DataFrame
import io,sys

from folder_finding import find_docx_in_folder
from docx.enum.table import WD_TABLE_DIRECTION
# For test 
#path= "C:/Users/ygkwo/Desktop/test1/1.docx" #input_path
#output_path="C:/Users/ygkwo/Desktop/test1" #output_path


# Refrence (완성 시 삭제)
# one file
#docx_to_xlsx Function [Input, Output]
# def docx_to_xlsx(input,output):
#     document=Document(path)

#     writer=pd.ExcelWriter('{}/docx_talbes.xlsx'.format(output),engine='xlsxwriter')

#     for i in range(len(document.tables)):
#         table=document.tables[i]
#         data=[]
#         keys=None
#         row_data=None
        
#         for j,row in enumerate(table.rows):
#             text=(cell.text for cell in row.cells)
#             if j==0:
#                 keys=tuple(text)
#                 continue
#             row_data=dict(zip(keys,text))
#             data.append(row_data)
#             #
#             # df1=pd.DataFrame(data)
#             # print(df1)
#             # df1.to_excel(writer,sheet_name='table_{}'.format(i))

#         df=pd.DataFrame(data)
#         print(df)
#         df.to_excel(writer,sheet_name='table_{}'.format(i))


#     writer.save()


#multi
def docx_to_xlsx_multi(input):

    for k in input:
        print("file is "+ k)
        print("\n")

        document=Document(k)

        # for another word file format
        if k.endswith(".docx") or k.endswith(".docm") or k.endswith(".dotm") or k.endswith(".dotx"): 
            out=k[:-5]

        elif k.endswith(".doc") or k.endswith(".dot"):
            out=k[:-4]

        #out이 slicing word file 경로의 확장자만 slicing 한 것

        #checkpoint
        print("check point: slicing path is"+out)

        
        #새로 붙이기
        writer=pd.ExcelWriter('{}.xlsx'.format(out),engine='xlsxwriter')

        for i in range(len(document.tables)):
            table=document.tables[i]
            data=[]
            keys=None
            row_data=None
            #print("This is talbe_direction")
            #print(table.direction)

            # 행과 열 수 비교하는 함수를.... 
            for j,row in enumerate(table.rows):
                text=(cell.text for cell in row.cells)
                if j==0:
                    keys=tuple(text)
                    continue
                row_data=dict(zip(keys,text))
                data.append(row_data)
                #
                # df1=pd.DataFrame(data)
                # print(df1)
                # df1.to_excel(writer,sheet_name='table_{}'.format(i))

            df=pd.DataFrame(data)
            print(df)
            df.to_excel(writer,sheet_name='table_{}'.format(i))

        writer.save()


#excel 다른 칸으로 붙이는 법! 


##완성시 삭제!!! 

#one docx file
#docx_to_xlsx(path,output_path)

# print("Input Folder's PATH")
# input_path=input()

# print("Input_PATH is ",input_path)

# #If you want to reflect what you want to input_path, modify parameter

# #find_docx_in_folder(input_path)

# in_1=find_docx_in_folder(input_path)
# print("We'll extract this files(.docx)")
# print(in_1)


# docx_to_xlsx_multi(in_1)
