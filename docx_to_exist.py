from docx.api import Document
from openpyxl import Workbook
from openpyxl import load_workbook
#from openpyxl import append_df_to_excel
import pandas as pd 
from pandas import DataFrame
import io,sys
import os
#import openpyxl

from folder_finding import find_docx_in_folder






#multi
def docx_to_xlsx_exist(input):
    #Input이 파일 이름배열이고 그 중에서 하나씩 뽑아낸 것이 k
    for k in input:
        print("file is "+ k)
        print("\n")

        document=Document(k)

        # for another word file format
        if k.endswith(".docx") or k.endswith(".docm") or k.endswith(".dotm") or k.endswith(".dotx"): 
            out=k[:-5]
        
        elif k.endswith(".doc") or k.endswith(".dot"):
            out=k[:-4]

        #out이 slicing word file 경로의 확장자만 slicing 한 것

        #checkpoint
        print("check point: slicing path is"+out)

        
        # if os.path.exists(out_xlsx):
        #     new_wb=load_workbook(out_xlsx)
        # else:
        #     new_wb=Workbook()


        #새로 붙이기
        writer=pd.ExcelWriter('{}.xlsx'.format(out),engine="openpyxl")
        #engine='xlsxwriter'

        for i in range(len(document.tables)):
            table=document.tables[i]
            data=[]
            keys=None
            row_data=None
            
            # 행과 열 수 비교하는 함수를.... 
            for j,row in enumerate(table.rows):
                text=(cell.text for cell in row.cells)
                if j==0:
                    keys=tuple(text)
                    continue
                row_data=dict(zip(keys,text))
                data.append(row_data)
                
            df=pd.DataFrame(data)
            print(df)
            #df.to_excel(writer,startrow= 5, startcol=5, sheet_name='table_{}'.format(i))
            append_df_to_excel(writer,df,sheet_name='table_{}'.format(i),startrow=5,startcol=5)
        writer.save()


#excel 다른 칸으로 붙이는 법! 
